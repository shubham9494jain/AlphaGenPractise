from rest_framework.decorators import api_view, permission_classes, action
from rest_framework.response import Response
from rest_framework import viewsets, status
from .models import Project, Document, SharedFile, ChatHistory
from .serializers import ProjectSerializer, DocumentSerializer, SharedFileSerializer, ChatHistorySerializer
from rest_framework.permissions import IsAuthenticated
import zipfile
import os
from django.conf import settings
from django.db.models import Q # Added Q import
import google.generativeai as genai # Import Gemini SDK
import io # For handling binary data
import base64 # For encoding/decoding binary data
from PIL import Image # For image processing (Pillow)
import pytesseract # For OCR
import PyPDF2 # For PDF text extraction
from docx import Document as DocxDocument # For DOCX text extraction
import pandas as pd # For CSV and XLSX

# Set the path to the Tesseract executable
# This is often needed on Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Configure Gemini API
genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))

def _extract_text_from_document(document):
    """
    Extracts text content from a Document object based on its content_type.
    Supports text, images (OCR), PDFs, DOCX, CSV, and XLSX.
    """
    if not document or not document.file_content:
        return ""

    content_type = document.content_type
    file_content_bytes = document.file_content

    try:
        if content_type == 'text/plain':
            try:
                return file_content_bytes.decode('utf-8')
            except UnicodeDecodeError:
                return file_content_bytes.decode('latin-1') # Fallback for other encodings
        elif content_type.startswith('image/'):
            # Use Pillow to open the image and pytesseract for OCR
            image = Image.open(io.BytesIO(file_content_bytes))
            text = pytesseract.image_to_string(image)
            return text
        elif content_type == 'application/pdf':
            # Use PyPDF2 to extract text from PDF
            reader = PyPDF2.PdfReader(io.BytesIO(file_content_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Use python-docx to extract text from DOCX
            doc = DocxDocument(io.BytesIO(file_content_bytes))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        elif content_type == 'text/csv':
            # Use pandas to read CSV and convert to string
            df = pd.read_csv(io.BytesIO(file_content_bytes))
            return df.to_string()
        elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            # Use pandas to read XLSX and convert to string
            df = pd.read_excel(io.BytesIO(file_content_bytes))
            return df.to_string()
        else:
            return f"Content of '{document.name}' (Unsupported content type: {content_type} for direct text extraction)."
    except Exception as e:
        return f"Error extracting text from '{document.name}' ({content_type}): {e}"

def get_ai_response(user_message, project=None, document=None):
    """
    Generates an AI response using Google Gemini, with a fallback mechanism.
    """
    model_pro = genai.GenerativeModel('gemini-2.5-pro')
    model_flash = genai.GenerativeModel('gemini-2.5-flash')

    document_context = ""
    if document:
        document_context = _extract_text_from_document(document)
        if document_context:
            document_context = f"\n\nDocument Context (from {document.name}):\n{document_context}\n"
    elif project:
        # For project-level chat, we could potentially combine content from all documents in the project
        # For now, we'll just indicate the project context.
        document_context = f"\n\nProject Context (from {project.title}). No specific document selected.\n"

    full_prompt = f"User query: {user_message}{document_context}\n\n"

    # Add specific instructions for predefined queries
    user_message_lower = user_message.lower()
    if "summarize this report" in user_message_lower:
        full_prompt += "Please provide a concise summary of the provided document/project context."
    elif "list top holdings" in user_message_lower:
        full_prompt += "Please list the top holdings or key entities mentioned in the provided document/project context."
    else:
        full_prompt += "Please answer the user's query based on the provided document/project context. If the query is unrelated to the context, state that you cannot help with it."

    try:
        # Try Gemini 1.5 Pro first
        response = model_pro.generate_content(full_prompt)
        return f"AI: {response.text}"
    except Exception as e:
        print(f"Gemini 1.5 Pro failed, trying Flash. Error: {e}")
        # Fallback to Gemini 1.5 Flash
        try:
            response = model_flash.generate_content(full_prompt)
            return f"AI: {response.text}"
        except Exception as flash_e:
            print(f"Gemini 1.5 Flash also failed. Error: {flash_e}")
            return "AI: I encountered an error trying to process your request with both Gemini Pro and Flash models. Please try again later or rephrase your query."

@api_view(['GET']) 
def home_page(request):
    data = {
        "message": "Hello, Lala! Ye React ke liye Django API se aa raha hai."
    }
    return Response(data)

from django.contrib.auth.models import User
from django.core.cache import cache
from django.core.mail import send_mail
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from django.utils.crypto import get_random_string
from .serializers import MyTokenObtainPairSerializer
from rest_framework_simplejwt.views import TokenObtainPairView
from google.oauth2 import id_token
from google.auth.transport import requests
from rest_framework_simplejwt.tokens import RefreshToken

class MyTokenObtainPairView(TokenObtainPairView):
    serializer_class = MyTokenObtainPairSerializer

from django.template.loader import render_to_string
from django.core.mail import EmailMultiAlternatives

@api_view(['POST'])
@permission_classes([AllowAny])
def register(request):
    email = request.data.get('email')
    password = request.data.get('password')
    name = request.data.get('name')

    if not email or not password or not name:
        return Response({'error': 'Please provide all fields'}, status=status.HTTP_400_BAD_REQUEST)

    if User.objects.filter(email=email).exists():
        return Response({'error': 'Email already exists'}, status=status.HTTP_400_BAD_REQUEST)

    user = User.objects.create_user(username=email, email=email, password=password, first_name=name, is_active=False)

    otp = get_random_string(length=6, allowed_chars='1234567890')
    cache.set(email, otp, timeout=600)

    subject = 'Your OTP for AlphaGen'
    from_email = 'shubham.mct265@gmail.com'
    to = [email]

    text_content = f'Your OTP is: {otp}'
    html_content = render_to_string('api/otp_email.html', {'otp': otp})

    msg = EmailMultiAlternatives(subject, text_content, from_email, to)
    msg.attach_alternative(html_content, "text/html")
    msg.send()

    return Response({'message': 'OTP sent to your email. Please verify.'}, status=status.HTTP_201_CREATED)

@api_view(['POST'])
@permission_classes([AllowAny])
def verify_otp(request):
    email = request.data.get('email')
    otp = request.data.get('otp')

    if not email or not otp:
        return Response({'error': 'Please provide both email and OTP'}, status=status.HTTP_400_BAD_REQUEST)

    cached_otp = cache.get(email)

    if not cached_otp:
        return Response({'error': 'OTP expired or invalid'}, status=status.HTTP_400_BAD_REQUEST)

    if otp != cached_otp:
        return Response({'error': 'Invalid OTP'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        user = User.objects.get(email=email)
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

    user.is_active = True
    user.save()
    cache.delete(email)

    return Response({'message': 'Account verified successfully. You can now login.'}, status=status.HTTP_200_OK)

@api_view(['POST'])
@permission_classes([AllowAny])
def resend_otp(request):
    email = request.data.get('email')
    if not email:
        return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        user = User.objects.get(email=email)
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

    if user.is_active:
        return Response({'message': 'This account is already active.'}, status=status.HTTP_400_BAD_REQUEST)

    otp = get_random_string(length=6, allowed_chars='1234567890')
    cache.set(email, otp, timeout=600)

    subject = 'Your New OTP for AlphaGen'
    from_email = 'shubham.mct265@gmail.com'
    to = [email]

    text_content = f'Your new OTP is: {otp}'
    html_content = render_to_string('api/otp_email.html', {'otp': otp})

    msg = EmailMultiAlternatives(subject, text_content, from_email, to)
    msg.attach_alternative(html_content, "text/html")
    msg.send()

    return Response({'message': 'A new OTP has been sent to your email.'}, status=status.HTTP_200_OK)

@api_view(['POST'])
@permission_classes([AllowAny])
def google_login(request):
    access_token = request.data.get('access_token')
    if not access_token:
        return Response({'error': 'Access token not provided'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        idinfo = id_token.verify_oauth2_token(access_token, requests.Request())
        email = idinfo['email']
        first_name = idinfo.get('given_name', '')
        last_name = idinfo.get('family_name', '')

        user, created = User.objects.get_or_create(
            email=email,
            defaults={
                'username': email,
                'first_name': first_name,
                'last_name': last_name,
                'is_active': True
            }
        )

        if created:
            user.set_unusable_password()
            user.save()

        refresh = RefreshToken.for_user(user)
        return Response({
            'refresh': str(refresh),
            'access': str(refresh.access_token),
        })

    except ValueError:
        return Response({'error': 'Invalid token'}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
@permission_classes([AllowAny])
def password_reset_request(request):
    email = request.data.get('email')
    if not email:
        return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        user = User.objects.get(email=email)
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

    otp = get_random_string(length=6, allowed_chars='1234567890')
    cache.set(f'password_reset_{email}', otp, timeout=600)

    subject = 'Your Password Reset OTP for AlphaGen'
    from_email = 'shubham.mct265@gmail.com'
    to = [email]

    text_content = f'Your password reset OTP is: {otp}'
    html_content = render_to_string('api/otp_email.html', {'otp': otp, 'title': 'Password Reset OTP'})

    msg = EmailMultiAlternatives(subject, text_content, from_email, to)
    msg.attach_alternative(html_content, "text/html")
    msg.send()

    return Response({'message': 'A password reset OTP has been sent to your email.'}, status=status.HTTP_200_OK)

@api_view(['POST'])
@permission_classes([AllowAny])
def password_reset_verify(request):
    email = request.data.get('email')
    otp = request.data.get('otp')

    if not email or not otp:
        return Response({'error': 'Please provide both email and OTP'}, status=status.HTTP_400_BAD_REQUEST)

    cached_otp = cache.get(f'password_reset_{email}')

    if not cached_otp:
        return Response({'error': 'OTP expired or invalid'}, status=status.HTTP_400_BAD_REQUEST)

    if otp != cached_otp:
        return Response({'error': 'Invalid OTP'}, status=status.HTTP_400_BAD_REQUEST)

    token = get_random_string(length=32)
    cache.set(f'password_reset_token_{email}', token, timeout=600)
    cache.delete(f'password_reset_{email}')

    return Response({'message': 'OTP verified successfully.', 'token': token}, status=status.HTTP_200_OK)

@api_view(['POST'])
@permission_classes([AllowAny])
def password_reset_confirm(request):
    email = request.data.get('email')
    token = request.data.get('token')
    new_password = request.data.get('new_password')

    if not email or not token or not new_password:
        return Response({'error': 'All fields are required'}, status=status.HTTP_400_BAD_REQUEST)

    reset_token = cache.get(f'password_reset_token_{email}')

    if not reset_token or token != reset_token:
        return Response({'error': 'Invalid or expired token'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        user = User.objects.get(email=email)
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

    user.set_password(new_password)
    user.save()
    cache.delete(f'password_reset_token_{email}')

    return Response({'message': 'Password has been reset successfully.'}, status=status.HTTP_200_OK)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def logout(request):
    try:
        refresh_token = request.data["refresh"]
        token = RefreshToken(refresh_token)
        token.blacklist()
        return Response(status=status.HTTP_205_RESET_CONTENT)
    except Exception as e:
        return Response(status=status.HTTP_400_BAD_REQUEST)

class ProjectViewSet(viewsets.ModelViewSet):
    serializer_class = ProjectSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Project.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    @action(detail=True, methods=['post'])
    def upload_document(self, request, pk=None):
        project = self.get_object()
        files = request.FILES.getlist('file')

        if not files:
            return Response({'error': 'No files provided'}, status=status.HTTP_400_BAD_REQUEST)

        for f in files:
            file_content = f.read()
            content_type = f.content_type
            
            document = Document(
                project=project,
                user=request.user,
                name=f.name,
                file_content=file_content,
                content_type=content_type
            )
            document.save()
        
        return Response({'message': 'Documents uploaded successfully'}, status=status.HTTP_201_CREATED)


class DocumentViewSet(viewsets.ModelViewSet):
    serializer_class = DocumentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        queryset = Document.objects.filter(user=self.request.user)
        project_id = self.request.query_params.get('project', None)
        if project_id is not None:
            queryset = queryset.filter(project__id=project_id)
        return queryset

    def perform_create(self, serializer):
        # This will be handled by the upload_document action in ProjectViewSet
        project_id = self.request.data.get('project')
        project = Project.objects.get(id=project_id, user=self.request.user)
        serializer.save(user=self.request.user, project=project)

class SharedFileViewSet(viewsets.ModelViewSet):
    serializer_class = SharedFileSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return SharedFile.objects.filter(user=self.request.user)

class ChatHistoryViewSet(viewsets.ModelViewSet):
    serializer_class = ChatHistorySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        queryset = ChatHistory.objects.filter(
            Q(project__user=self.request.user) | Q(document__user=self.request.user)
        )
        project_id = self.request.query_params.get('project_id')
        document_id = self.request.query_params.get('document_id')

        if project_id:
            queryset = queryset.filter(project__id=project_id)
        if document_id:
            queryset = queryset.filter(document__id=document_id)
        
        return queryset.order_by('created_at')

    def create(self, request, *args, **kwargs):
        user_message = request.data.get('message')
        project_id = request.data.get('project_id')
        document_id = request.data.get('document_id')
        
        if not user_message:
            return Response({'error': 'Message is required'}, status=status.HTTP_400_BAD_REQUEST)

        if not project_id and not document_id:
            return Response({'error': 'Either project_id or document_id is required'}, status=status.HTTP_400_BAD_REQUEST)

        project = None
        document = None

        if project_id:
            try:
                project = Project.objects.get(id=project_id, user=request.user)
            except Project.DoesNotExist:
                return Response({'error': 'Project not found'}, status=status.HTTP_404_NOT_FOUND)
        
        if document_id:
            try:
                document = Document.objects.get(id=document_id, user=request.user)
            except Document.DoesNotExist:
                return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Save user message
        user_chat_entry = ChatHistory.objects.create(
            project=project,
            document=document,
            message=user_message,
            is_user_message=True
        )
        user_serializer = self.get_serializer(user_chat_entry)

        # Get AI response
        ai_response_text = get_ai_response(user_message, project, document)

        # Save AI response
        ai_chat_entry = ChatHistory.objects.create(
            project=project,
            document=document,
            message=ai_response_text,
            is_user_message=False
        )
        ai_serializer = self.get_serializer(ai_chat_entry)

        return Response({
            'user_message': user_serializer.data,
            'ai_message': ai_serializer.data
        }, status=status.HTTP_201_CREATED)